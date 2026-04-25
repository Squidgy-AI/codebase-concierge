"""
File upload pipeline. Converts everything to PDF in-memory and pushes to Nia.

Three-step Nia upload:
  1. POST /v2/sources/upload-url        → signed GCS URL + gcs_path
  2. PUT bytes to that signed URL       → file lands in GCS
  3. POST /v2/sources type=pdf gcs_path → indexing kicks off

For non-PDFs we convert text→PDF in-memory so the operator sees one consistent
"Indexed sources" surface regardless of original format.
"""
import io
import os

import httpx


NIA_BASE = "https://apigcp.trynia.ai"
NIA_API_KEY = os.environ.get("NIA_API_KEY", "").strip()


# ---------- Conversion helpers ----------

def text_to_pdf_bytes(text: str, title: str = "Document") -> bytes:
    """Render plain text as a minimal multi-page PDF. Latin-1 safe."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(left=15, top=15, right=15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    # fpdf2 emits Latin-1 by default; replace anything outside that range so we
    # don't crash on stray smart quotes / em-dashes / emoji.
    safe = text.encode("latin-1", "replace").decode("latin-1")
    width = pdf.w - pdf.l_margin - pdf.r_margin
    for line in safe.splitlines() or [""]:
        if not line.strip():
            pdf.ln(4)
            continue
        pdf.multi_cell(width, 5, line, new_x="LMARGIN", new_y="NEXT")
    out = pdf.output(dest="S")
    return bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")


def docx_bytes_to_text(content: bytes) -> str:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Tables flatten to tab-separated rows.
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n\n".join(parts)


def normalize_to_pdf(filename: str, content: bytes) -> tuple[bytes, str]:
    """Return (pdf_bytes, suggested_display_name) for any supported file."""
    ext = (os.path.splitext(filename or "")[1] or "").lower().lstrip(".")
    stem = os.path.splitext(os.path.basename(filename or "Document"))[0] or "Document"

    if ext == "pdf":
        return content, stem
    if ext in ("txt", "md"):
        text = content.decode("utf-8", errors="replace")
        return text_to_pdf_bytes(text, title=stem), stem
    if ext == "docx":
        text = docx_bytes_to_text(content)
        return text_to_pdf_bytes(text, title=stem), stem
    raise ValueError(f"unsupported file type: .{ext}")


# ---------- Nia upload ----------

async def _signed_upload_url(filename: str, content_type: str = "application/pdf") -> dict:
    async with httpx.AsyncClient(timeout=15) as c:
        r = await c.post(
            f"{NIA_BASE}/v2/sources/upload-url",
            headers={"Authorization": f"Bearer {NIA_API_KEY}"},
            json={"filename": filename, "content_type": content_type},
        )
        r.raise_for_status()
        return r.json()


async def _put_to_gcs(upload_url: str, body: bytes, content_type: str = "application/pdf") -> None:
    async with httpx.AsyncClient(timeout=120) as c:
        r = await c.put(upload_url, content=body, headers={"Content-Type": content_type})
        r.raise_for_status()


async def _create_pdf_source(gcs_path: str, display_name: str | None) -> dict:
    # Nia exposes uploaded PDFs as type=research_paper with is_pdf=true.
    # Per Nia: gcs_path is only valid on type=documentation; is_pdf flag tells the
    # ingester this is a binary PDF rather than a crawl URL.
    payload: dict = {"type": "documentation", "is_pdf": True, "gcs_path": gcs_path}
    async with httpx.AsyncClient(timeout=30) as c:
        r = await c.post(
            f"{NIA_BASE}/v2/sources",
            headers={"Authorization": f"Bearer {NIA_API_KEY}"},
            json=payload,
        )
        if r.status_code >= 400:
            print(f"[upload] /v2/sources rejected: {r.status_code} {r.text}")
        r.raise_for_status()
        src = r.json()
        if display_name:
            try:
                pr = await c.patch(
                    f"{NIA_BASE}/v2/sources/{src['id']}",
                    headers={"Authorization": f"Bearer {NIA_API_KEY}"},
                    json={"display_name": display_name},
                )
                pr.raise_for_status()
                src["display_name"] = display_name
            except httpx.HTTPError:
                pass
        return src


async def upload_file(filename: str, content: bytes, display_name: str | None = None) -> dict:
    """End-to-end: convert if needed → push to GCS → register source in Nia.
    Returns the Nia source row."""
    pdf_bytes, default_name = normalize_to_pdf(filename, content)
    name = (display_name or default_name).strip() or default_name

    signed = await _signed_upload_url(filename=os.path.basename(filename) or "upload.pdf")
    upload_url = signed.get("upload_url") or signed.get("url")
    gcs_path = signed.get("gcs_path") or signed.get("path") or signed.get("object")
    if not upload_url or not gcs_path:
        raise RuntimeError(f"unexpected Nia signed-url response: {signed}")

    await _put_to_gcs(upload_url, pdf_bytes, content_type="application/pdf")
    return await _create_pdf_source(gcs_path, display_name=name)
